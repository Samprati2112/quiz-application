import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEACHER_CODE = "TEACH2024"

def parse_docx(filepath):
    try:
        doc = Document(filepath)
    except Exception as exc:
        return [], [f"Cannot open file: {exc}"]

    raw_lines = [p.text.strip() for p in doc.paragraphs]

    blocks, current = [], []
    for line in raw_lines:
        if line == "" or re.match(r"^[-─=*]{3,}$", line):
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)

    def has_choices(block):
        return any(re.match(r"^[A-Da-d][\)\.]\s*.+", ln) for ln in block)

    question_blocks = [b for b in blocks if has_choices(b)]

    questions, errors = [], []
    for i, block in enumerate(question_blocks, 1):
        result, err = _parse_block(block)
        if err:
            errors.append(f"Question {i}: {err}")
        else:
            questions.append(result)

    if not questions and not errors:
        errors.append("No question blocks found. Check the file format.")

    return questions, errors


def generate_template(filepath):
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("QuizMaster Pro — Question Template")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    doc.add_paragraph()

    inst = doc.add_paragraph()
    inst.add_run("Instructions\n").bold = True
    inst.add_run(
        "• One question per block, separated by a blank line.\n"
        "• Q: prefix is optional — the parser picks up the first line automatically.\n"
        "• A) through D) are the four choices (you can use a) or A. etc.).\n"
        "• Correct: must be A, B, C, or D (the letter of the right choice).\n"
        "• Category and Difficulty are optional (defaults: General Knowledge / Medium).\n"
        "• Supported Difficulty values: Easy  |  Medium  |  Hard\n"
        "• Delete this instruction block before uploading.\n"
    )

    doc.add_paragraph()
    _divider(doc)
    doc.add_paragraph()

    examples = [
        ("What is the capital of France?",
         ["Paris", "London", "Berlin", "Madrid"], "A", "Geography", "Easy"),
        ("Who discovered penicillin?",
         ["Alexander Fleming", "Louis Pasteur", "Jonas Salk", "Joseph Lister"],
         "A", "Science", "Medium"),
    ]
    for q, choices, correct, cat, diff in examples:
        _write_block(doc, q, choices, correct, cat, diff)
        doc.add_paragraph()
        _divider(doc)
        doc.add_paragraph()

    for _ in range(5):
        _write_blank_block(doc)
        doc.add_paragraph()
        _divider(doc)
        doc.add_paragraph()

    doc.save(filepath)


_Q_RE      = re.compile(r"^(?:Q\s*[\.:]\s*|\d+[\.)]\s*)", re.IGNORECASE)
_CHOICE_RE = re.compile(r"^([A-Da-d])[\)\.]\s*(.+)")
_CORRECT_RE = re.compile(r"^(?:correct|answer)\s*[:=]\s*([A-Da-d])", re.IGNORECASE)
_CAT_RE    = re.compile(r"^category\s*[:=]\s*(.+)", re.IGNORECASE)
_DIFF_RE   = re.compile(r"^difficulty\s*[:=]\s*(.+)", re.IGNORECASE)


def _parse_block(lines):
    q_parts      = []
    choices      = {}          # "A" → text
    correct_key  = None
    category     = "General Knowledge"
    difficulty   = "Medium"
    reading_q    = False

    for line in lines:
        if _Q_RE.match(line):
            reading_q = True
            text = _Q_RE.sub("", line).strip()
            if text:
                q_parts.append(text)

        elif _CHOICE_RE.match(line):
            reading_q = False
            m = _CHOICE_RE.match(line)
            choices[m.group(1).upper()] = m.group(2).strip()

        elif _CORRECT_RE.match(line):
            reading_q = False
            correct_key = _CORRECT_RE.match(line).group(1).upper()

        elif _CAT_RE.match(line):
            reading_q = False
            category = _CAT_RE.match(line).group(1).strip().title()

        elif _DIFF_RE.match(line):
            reading_q = False
            raw = _DIFF_RE.match(line).group(1).strip().capitalize()
            difficulty = raw if raw in ("Easy", "Medium", "Hard") else "Medium"

        elif reading_q and line:
            q_parts.append(line)

    if not q_parts:
        for line in lines:
            if not (_CHOICE_RE.match(line) or _CORRECT_RE.match(line)
                    or _CAT_RE.match(line) or _DIFF_RE.match(line)):
                q_parts.append(line)
                break

    if not q_parts:
        return None, "No question text found."
    if len(choices) < 2:
        return None, f"Need at least 2 choices, found {len(choices)}."
    if not correct_key:
        return None, "Missing 'Correct: X' line."
    if correct_key not in choices:
        return None, f"'Correct: {correct_key}' but choice {correct_key} is missing."

    ordered_keys   = sorted(choices.keys())          # A, B, C, D
    ordered_choices = [choices[k] for k in ordered_keys]
    correct_idx    = ordered_keys.index(correct_key)

    return {
        "question":      " ".join(q_parts),
        "choices":       ordered_choices,
        "correct_index": correct_idx,
        "category":      category,
        "difficulty":    difficulty,
    }, None


def _write_block(doc, question, choices, correct, category, difficulty):
    _styled(doc, f"Q: {question}", bold=True)
    for letter, ch in zip("ABCD", choices):
        color = RGBColor(0x22, 0xC5, 0x5E) if letter == correct else None
        _styled(doc, f"{letter}) {ch}", color=color)
    _styled(doc, f"Correct: {correct}", color=RGBColor(0x22, 0xC5, 0x5E))
    _styled(doc, f"Category: {category}", color=RGBColor(0x64, 0x74, 0x8B))
    _styled(doc, f"Difficulty: {difficulty}", color=RGBColor(0x64, 0x74, 0x8B))


def _write_blank_block(doc):
    _styled(doc, "Q: ", bold=True, placeholder=True)
    for letter in "ABCD":
        _styled(doc, f"{letter}) ", placeholder=True)
    _styled(doc, "Correct: ", placeholder=True)
    _styled(doc, "Category: Geography", color=RGBColor(0x64, 0x74, 0x8B))
    _styled(doc, "Difficulty: Easy", color=RGBColor(0x64, 0x74, 0x8B))


def _styled(doc, text, bold=False, color=None, placeholder=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if placeholder:
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    elif color:
        run.font.color.rgb = color


def _divider(doc):
    p = doc.add_paragraph("─" * 72)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x66)
